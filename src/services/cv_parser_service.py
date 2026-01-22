"""
CV Parser Service

Extracts text content from CV files (PDF, DOCX, TXT, MD, TEX).
"""

import re
from pathlib import Path
from typing import Optional, Dict, Tuple
from loguru import logger


class CVParserService:
    """
    Service for extracting text from CV files.
    
    Supports:
    - PDF files
    - DOCX files
    - TXT files
    - Markdown files
    - LaTeX files (.tex)
    """
    
    def __init__(self):
        """Initialize the CV parser service."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check available parsing libraries."""
        self._has_pypdf = False
        self._has_docx = False
        
        try:
            import pypdf
            self._has_pypdf = True
        except ImportError:
            try:
                import PyPDF2
                self._has_pypdf = True
            except ImportError:
                logger.warning("pypdf/PyPDF2 not installed - PDF parsing limited")
        
        try:
            import docx
            self._has_docx = True
        except ImportError:
            logger.warning("python-docx not installed - DOCX parsing unavailable")
    
    def extract_text(self, file_path: Path) -> Optional[str]:
        """
        Extract text from a CV file.
        
        Args:
            file_path: Path to the CV file
            
        Returns:
            Extracted text or None if extraction failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        suffix = file_path.suffix.lower()
        
        extractors = {
            ".pdf": self._extract_from_pdf,
            ".docx": self._extract_from_docx,
            ".doc": self._extract_from_docx,
            ".txt": self._extract_from_text,
            ".md": self._extract_from_text,
            ".tex": self._extract_from_latex,
        }
        
        extractor = extractors.get(suffix)
        if not extractor:
            logger.error(f"Unsupported file format: {suffix}")
            return None
        
        try:
            text = extractor(file_path)
            if text:
                # Clean up text (but keep LaTeX structure for .tex files)
                if suffix != ".tex":
                    text = self._clean_text(text)
                return text
        except Exception as e:
            logger.error(f"Extraction failed: {e}")
        
        return None
    
    def extract_latex_with_structure(self, file_path: Path) -> Tuple[str, str, Dict]:
        """
        Extract LaTeX file preserving structure and detecting language.
        
        Returns:
            Tuple of (raw_content, extracted_text, structure_info)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return "", "", {}
        
        try:
            content = self._read_file(file_path)
            if not content:
                return "", "", {}
            
            # Detect language from LaTeX
            language = self._detect_latex_language(content)
            
            # Extract readable text
            text = self._latex_to_text(content)
            
            # Extract structure info
            structure = self._extract_latex_structure(content)
            structure["language"] = language
            structure["raw_latex"] = content
            
            return content, text, structure
            
        except Exception as e:
            logger.error(f"LaTeX extraction failed: {e}")
            return "", "", {}
    
    def _extract_from_latex(self, file_path: Path) -> Optional[str]:
        """Extract text from LaTeX file."""
        content = self._read_file(file_path)
        if not content:
            return None
        
        return self._latex_to_text(content)
    
    def _latex_to_text(self, latex_content: str) -> str:
        """Convert LaTeX to readable text."""
        text = latex_content
        
        # Remove comments
        text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)
        
        # Remove document class and packages
        text = re.sub(r'\\documentclass\[?[^\]]*\]?\{[^}]*\}', '', text)
        text = re.sub(r'\\usepackage\[?[^\]]*\]?\{[^}]*\}', '', text)
        
        # Extract content between begin/end document
        doc_match = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', text, re.DOTALL)
        if doc_match:
            text = doc_match.group(1)
        
        # Convert sections to readable format
        text = re.sub(r'\\section\*?\{([^}]*)\}', r'\n\n## \1\n', text)
        text = re.sub(r'\\subsection\*?\{([^}]*)\}', r'\n### \1\n', text)
        text = re.sub(r'\\subsubsection\*?\{([^}]*)\}', r'\n#### \1\n', text)
        
        # Convert text formatting
        text = re.sub(r'\\textbf\{([^}]*)\}', r'**\1**', text)
        text = re.sub(r'\\textit\{([^}]*)\}', r'*\1*', text)
        text = re.sub(r'\\emph\{([^}]*)\}', r'*\1*', text)
        text = re.sub(r'\\underline\{([^}]*)\}', r'\1', text)
        
        # Convert lists
        text = re.sub(r'\\begin\{itemize\}', '', text)
        text = re.sub(r'\\end\{itemize\}', '', text)
        text = re.sub(r'\\begin\{enumerate\}', '', text)
        text = re.sub(r'\\end\{enumerate\}', '', text)
        text = re.sub(r'\\item\s*', '\n- ', text)
        
        # Convert href
        text = re.sub(r'\\href\{([^}]*)\}\{([^}]*)\}', r'\2 (\1)', text)
        
        # Remove common LaTeX commands but keep content
        text = re.sub(r'\\[a-zA-Z]+\*?\{([^}]*)\}', r'\1', text)
        
        # Remove remaining LaTeX commands
        text = re.sub(r'\\[a-zA-Z]+\*?', '', text)
        
        # Remove braces
        text = re.sub(r'[{}]', '', text)
        
        # Clean up whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        return text.strip()
    
    def _detect_latex_language(self, content: str) -> str:
        """Detect language from LaTeX content."""
        # Check babel package
        babel_match = re.search(r'\\usepackage\[([^\]]*)\]\{babel\}', content)
        if babel_match:
            lang = babel_match.group(1).lower()
            if 'french' in lang or 'francais' in lang:
                return 'fr'
            if 'english' in lang:
                return 'en'
        
        # Check for French-specific content
        french_indicators = [
            r'Expérience', r'Compétences', r'Formation', r'Langues',
            r'Résumé', r'Coordonnées', r'Professionnel'
        ]
        for indicator in french_indicators:
            if re.search(indicator, content, re.IGNORECASE):
                return 'fr'
        
        return 'en'
    
    def _extract_latex_structure(self, content: str) -> Dict:
        """Extract structure information from LaTeX."""
        structure = {
            "sections": [],
            "has_photo": False,
            "document_class": "",
            "packages": []
        }
        
        # Document class
        class_match = re.search(r'\\documentclass\[?([^\]]*)\]?\{([^}]*)\}', content)
        if class_match:
            structure["document_class"] = class_match.group(2)
        
        # Packages
        packages = re.findall(r'\\usepackage\[?[^\]]*\]?\{([^}]*)\}', content)
        structure["packages"] = packages
        
        # Sections
        sections = re.findall(r'\\section\*?\{([^}]*)\}', content)
        structure["sections"] = sections
        
        # Check for photo
        if re.search(r'\\includegraphics|\\photo', content):
            structure["has_photo"] = True
        
        return structure
    
    def _read_file(self, file_path: Path) -> Optional[str]:
        """Read file with encoding detection."""
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        return None
    
    def _extract_from_pdf(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF file."""
        if not self._has_pypdf:
            return self._extract_pdf_command(file_path)
        
        try:
            try:
                from pypdf import PdfReader
            except ImportError:
                from PyPDF2 import PdfReader
            
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            if text_parts:
                return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}")
        
        return self._extract_pdf_command(file_path)
    
    def _extract_pdf_command(self, file_path: Path) -> Optional[str]:
        """Extract PDF using pdftotext command."""
        import subprocess
        
        try:
            result = subprocess.run(
                ["pdftotext", "-layout", str(file_path), "-"],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode == 0 and result.stdout:
                return result.stdout
                
        except FileNotFoundError:
            logger.warning("pdftotext not found - install poppler-utils")
        except Exception as e:
            logger.warning(f"pdftotext failed: {e}")
        
        return None
    
    def _extract_from_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX file."""
        if not self._has_docx:
            logger.error("python-docx not installed")
            return None
        
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return None
    
    def _extract_from_text(self, file_path: Path) -> Optional[str]:
        """Extract text from TXT or MD file."""
        return self._read_file(file_path)
    
    def _clean_text(self, text: str) -> str:
        """Clean up extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove page numbers
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        text = re.sub(r'Page \d+ of \d+', '', text)
        
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats."""
        formats = [".txt", ".md", ".tex"]
        
        if self._has_pypdf:
            formats.append(".pdf")
        
        if self._has_docx:
            formats.extend([".docx", ".doc"])
        
        return formats
